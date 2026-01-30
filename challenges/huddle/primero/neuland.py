from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_budget_docx():
    doc = Document()
    
    # Título
    title = doc.add_heading('Presupuesto Mano de Obra - Neuland', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Datos
    data = {
        "1. Cocina": [
            ("Colocación de termocalefón y conexión de agua corriente", "350.000")
        ],
        "2. Lavandería": [
            ("Colocación de machimbre con nivelación 12,25m²", "1.500.000"),
            ("Revoque lavandería 37,44m²", "1.800.000"),
            ("Pintura pared 37,44m²", "650.000"),
            ("Pintura cielo raso 12,25m²", "300.000"),
            ("Colocación de nuevo desagüe para lavarropas y salida de aire", "800.000"),
            ("Nueva conexión agua embutida con 1 salida", "1.200.000"),
            ("Colocación de 5 cajitas para enchufes embutida", "250.000"),
            ("Cortar pared y colocar 1 puerta", "750.000"),
            ("Colocación de cerámica 12,25m²", "950.000"),
            ("Colocación de zócalo 12m", "480.000")
        ],
        "3. Pieza lado derecho": [
            ("Desmontar 1 lavamanos 1 wáter - anular conexión de agua", "550.000")
        ],
        "4. Baño lado derecho": [
            ("Desmontar 5m² de pared y 1 puerta", "850.000"),
            ("Nuevo cimiento encadenado 2,40m", "400.000"),
            ("Nueva pared divisoria 6,72m²", "650.000"),
            ("Colocación de 1 puerta", "250.000"),
            ("Colocación de cielo raso con nivelación 11,76m²", "1.300.000"),
            ("Pintura cielo raso 11,76m²", "280.000"),
            ("Colocación de desagüe (lavamanos-water-ducha) a pozo", "2.500.000"),
            ("Plomería de agua y termocalefón", "1.900.000"),
            ("Colocación de azulejos 28m²", "3.500.000"),
            ("Colocación de Piso cerámica 11,76m²", "900.000"),
            ("Colocación de termocalefón - water - lavamanos", "750.000"),
            ("Revoque 18,76m²", "880.000"),
            ("Pintura 18,76m²", "300.000"),
            ("Cámara séptica", "950.000")
        ],
        "5. Sala": [
            ("Revoque 57m²", "2.700.000"),
            ("Pintura 57m²", "900.000"),
            ("Colocación de 7 cajitas para luz", "350.000")
        ],
        "6. Baño lado izquierdo": [
            ("Colocación de 1 llave mono comando + termocalefón", "1.200.000"),
            ("Desmontar azulejos 6,48m²", "350.000"),
            ("Colocación de nuevo azulejos 6,48m²", "800.000"),
            ("Divisoria para baño 1,50x2 de altura", "600.000")
        ],
        "7. Varios": [
            ("Reparación de 9 ventanas", "400.000")
        ]
    }

    grand_total = 0
    
    # Generar tablas
    for section, items in data.items():
        doc.add_heading(section, level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Descripción'
        hdr_cells[1].text = 'Costo (Gs.)'
        
        section_total = 0
        
        for desc, cost_str in items:
            row_cells = table.add_row().cells
            row_cells[0].text = desc
            row_cells[1].text = cost_str
            # Calcular total (limpiando puntos)
            cost_int = int(cost_str.replace('.', ''))
            section_total += cost_int

        # Fila de Subtotal
        row_cells = table.add_row().cells
        row_cells[0].text =